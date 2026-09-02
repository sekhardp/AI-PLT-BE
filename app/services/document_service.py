import asyncio
import io
import logging
import uuid
from typing import Any, List, Optional
from google import genai
from pypdf import PdfReader
import docx
from sqlalchemy import delete, func, select, or_
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.settings import app_settings
from app.db.rag_models import DocumentChunk, UserDocument
from app.services.user_service import user_service

logger = logging.getLogger(__name__)

# User Quota Limits
MAX_DOCS_PER_USER = 5
MAX_STORAGE_BYTES = 100 * 1024 * 1024  # 100 MB


class DocumentService:
    """
    Document service managing parsing (PDF/DOCX/TXT), semantic chunking,
    Vertex AI vectorization (embeddings), user storage quotas, and persistence.
    """

    def __init__(self):
        self.gcp_project = "beam-suntory-gemini-llm-poc"
        self.gcp_location = "us-central1"
        self.embedding_model = "text-embedding-005"
        self._genai_client = None

    @property
    def genai_client(self):
        if self._genai_client is None:
            try:
                self._genai_client = genai.Client(
                    vertexai=True,
                    project=self.gcp_project,
                    location=self.gcp_location,
                )
            except Exception as e:
                logger.warning("Failed to initialize Google GenAI Client: %s", e)
                return None
        return self._genai_client

    async def _resolve_user_identifiers(self, user_id: str | None, db: AsyncSession) -> tuple[bool, List[str]]:
        """Resolve admin status and list of user identifiers for querying."""
        if not user_id:
            return False, []
        is_admin = await user_service.is_admin_user(user_id, db)
        if is_admin:
            return True, []
        user = await user_service.resolve_user(user_id, db)
        identifiers = {str(user_id).strip()}
        if user:
            identifiers.add(str(user.id))
            if user.email:
                identifiers.add(user.email.lower().strip())
        return False, list(identifiers)

    async def validate_quota(self, user_id: str, new_file_size: int, db: AsyncSession) -> None:
        """Enforce max 5 documents and max 100MB cumulative storage per user."""
        is_admin, identifiers = await self._resolve_user_identifiers(user_id, db)
        if is_admin:
            return

        cond = UserDocument.user_id.in_(identifiers) if identifiers else (UserDocument.user_id == user_id)

        # 1. Check document count
        count = await db.scalar(
            select(func.count(UserDocument.id)).where(cond)
        )
        if (count or 0) >= MAX_DOCS_PER_USER:
            raise ValueError(
                f"Document limit reached. You can upload at most {MAX_DOCS_PER_USER} documents."
            )

        # 2. Check cumulative storage
        total_bytes = await db.scalar(
            select(func.coalesce(func.sum(UserDocument.file_size_bytes), 0)).where(cond)
        )
        total_bytes = total_bytes or 0
        if (total_bytes + new_file_size) > MAX_STORAGE_BYTES:
            remaining_mb = max(0.0, (MAX_STORAGE_BYTES - total_bytes) / (1024 * 1024))
            raise ValueError(
                f"Storage quota exceeded. You have {remaining_mb:.1f} MB remaining out of 100 MB."
            )

    def extract_text(self, filename: str, content: bytes, mime_type: str = "") -> str:
        """Extract clean text from PDF, DOCX, or text/markdown formats."""
        lower_name = filename.lower()
        if lower_name.endswith(".pdf") or "pdf" in mime_type:
            reader = PdfReader(io.BytesIO(content))
            pages_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(text)
            return "\n\n".join(pages_text)

        if lower_name.endswith(".docx") or "word" in mime_type:
            doc = docx.Document(io.BytesIO(content))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

        # Default: treat as utf-8 plain text / markdown / code
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1", errors="replace")

    def chunk_text(
        self,
        text: str,
        chunk_size: int = 800,
        overlap: int = 100,
        separators: Optional[List[str]] = None,
    ) -> List[str]:
        """
        Split raw text into overlapping semantic passages using a recursive character splitter.
        Splits hierarchically by paragraphs, lines, sentences, and words to preserve context.
        """
        clean_text = text.strip()
        if not clean_text:
            return []

        if len(clean_text) <= chunk_size:
            return [clean_text]

        if separators is None:
            separators = ["\n\n", "\n", ". ", "? ", "! ", " ", ""]

        def _split_text(text_to_split: str, seps: List[str]) -> List[str]:
            final_chunks: List[str] = []
            if not seps or not text_to_split:
                return [text_to_split] if text_to_split else []

            sep = seps[0]
            new_seps = seps[1:]

            if sep == "":
                splits = list(text_to_split)
                separator_used = ""
            else:
                splits = text_to_split.split(sep)
                separator_used = sep

            good_splits: List[str] = []
            for s in splits:
                if not s and sep != "":
                    continue
                if len(s) < chunk_size:
                    good_splits.append(s)
                else:
                    if good_splits:
                        merged = _merge_splits(good_splits, separator_used)
                        final_chunks.extend(merged)
                        good_splits = []
                    if new_seps:
                        sub_chunks = _split_text(s, new_seps)
                        final_chunks.extend(sub_chunks)
                    else:
                        final_chunks.append(s)

            if good_splits:
                merged = _merge_splits(good_splits, separator_used)
                final_chunks.extend(merged)

            return [c.strip() for c in final_chunks if c.strip()]

        def _merge_splits(splits: List[str], separator: str) -> List[str]:
            docs: List[str] = []
            current_doc: List[str] = []
            total_len = 0

            for d in splits:
                d_len = len(d) + (len(separator) if current_doc else 0)
                if total_len + d_len > chunk_size and current_doc:
                    doc_text = separator.join(current_doc).strip()
                    if doc_text:
                        docs.append(doc_text)
                    while current_doc and (total_len > overlap or total_len + d_len > chunk_size):
                        popped = current_doc.pop(0)
                        total_len -= len(popped) + (len(separator) if current_doc else 0)

                current_doc.append(d)
                total_len += d_len

            if current_doc:
                doc_text = separator.join(current_doc).strip()
                if doc_text:
                    docs.append(doc_text)

            return docs

        return _split_text(clean_text, separators)

    async def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate 768-dim embeddings in batches using Vertex AI text-embedding-005."""
        if not texts:
            return []

        client = self.genai_client
        if not client:
            return []

        embeddings = []
        batch_size = 10
        for i in range(0, len(texts), batch_size):
            batch = texts[i : i + batch_size]
            try:
                response = await asyncio.wait_for(
                    client.aio.models.embed_content(
                        model=self.embedding_model,
                        contents=batch,
                    ),
                    timeout=5.0,
                )
                if hasattr(response, "embeddings"):
                    for emb in response.embeddings:
                        embeddings.append(emb.values)
                elif hasattr(response, "embedding"):
                    embeddings.append(response.embedding.values)
            except Exception as e:
                logger.warning("Embedding generation timed out or failed: %s", e)
                break

        return embeddings

    async def ingest_document(
        self,
        user_id: str,
        filename: str,
        content: bytes,
        mime_type: str,
        db: AsyncSession,
    ) -> UserDocument:
        """Parse, validate quota, chunk, embed, and store document in Cloud SQL."""
        file_size = len(content)
        await self.validate_quota(user_id, file_size, db)

        # Resolve to canonical numeric user ID if user exists
        user = await user_service.resolve_user(user_id, db)
        stored_user_id = str(user.id) if user else (str(user_id).strip() if user_id else "default_user")

        # 1. Create document record with 'indexing' status
        doc = UserDocument(
            user_id=stored_user_id,
            filename=filename,
            file_size_bytes=file_size,
            mime_type=mime_type or "application/octet-stream",
            status="indexing",
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)

        try:
            # 2. Extract and chunk text
            raw_text = self.extract_text(filename, content, mime_type)
            if not raw_text.strip():
                raise ValueError("No extractable text found in uploaded document.")

            chunks = self.chunk_text(raw_text)
            if not chunks:
                raise ValueError("Document yielded 0 text chunks.")

            # 3. Generate embeddings
            vectors = await self.generate_embeddings(chunks)

            # 4. Insert chunks into document_chunks
            for idx, (chunk_text, vec) in enumerate(zip(chunks, vectors)):
                chunk_record = DocumentChunk(
                    document_id=doc.id,
                    chunk_index=idx,
                    chunk_text=chunk_text,
                    token_count=len(chunk_text.split()),
                    embedding=vec,
                )
                db.add(chunk_record)

            doc.status = "ready"
            await db.commit()
            await db.refresh(doc)
            return doc
        except Exception as e:
            logger.error("Failed to ingest document %s: %s", filename, e)
            doc.status = "failed"
            doc.error_message = str(e)
            await db.commit()
            raise

    async def list_documents(self, user_id: str, db: AsyncSession) -> dict[str, Any]:
        """List user documents and compute quota metrics."""
        query = select(UserDocument).order_by(UserDocument.created_at.desc())
        is_admin, identifiers = await self._resolve_user_identifiers(user_id, db)
        if identifiers and not is_admin:
            query = query.where(UserDocument.user_id.in_(identifiers))
        result = await db.scalars(query)
        docs = result.all()

        total_bytes = sum(d.file_size_bytes for d in docs)
        return {
            "documents": [
                {
                    "id": str(d.id),
                    "filename": d.filename,
                    "file_size_bytes": d.file_size_bytes,
                    "file_size_mb": round(d.file_size_bytes / (1024 * 1024), 2),
                    "mime_type": d.mime_type,
                    "status": d.status,
                    "error_message": d.error_message,
                    "created_at": d.created_at.isoformat() if d.created_at else None,
                }
                for d in docs
            ],
            "quota": {
                "total_documents": len(docs),
                "max_documents": MAX_DOCS_PER_USER,
                "total_bytes": total_bytes,
                "total_mb": round(total_bytes / (1024 * 1024), 2),
                "max_mb": 100.0,
                "remaining_mb": round(max(0.0, (MAX_STORAGE_BYTES - total_bytes) / (1024 * 1024)), 2),
            },
        }

    async def delete_document(self, doc_id: str, user_id: str, db: AsyncSession) -> bool:
        """Delete a document and cascade its vectorized chunks."""
        doc_uuid = uuid.UUID(doc_id)
        query = select(UserDocument).where(UserDocument.id == doc_uuid)
        is_admin, identifiers = await self._resolve_user_identifiers(user_id, db)
        if identifiers and not is_admin:
            query = query.where(UserDocument.user_id.in_(identifiers))
        doc = await db.scalar(query)
        if not doc:
            return False

        await db.execute(delete(DocumentChunk).where(DocumentChunk.document_id == doc_uuid))
        await db.delete(doc)
        await db.commit()
        return True

document_service = DocumentService()


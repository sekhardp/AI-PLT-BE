import io
import logging
import uuid
from typing import Any, List, Optional
from google import genai
from pypdf import PdfReader
import docx
from sqlalchemy import delete, func, select, or_
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.settings import app_settings
from app.db.rag_models import DocumentChunk, UserDocument

logger = logging.getLogger(__name__)

# User Quota Limits
MAX_DOCS_PER_USER = 5
MAX_STORAGE_BYTES = 100 * 1024 * 1024  # 100 MB


class RAGService:
    """
    RAG service managing document parsing, chunking, Vertex AI embeddings,
    user storage quotas, and pgvector cosine similarity search.
    """

    def __init__(self):
        self.gcp_project = "beam-suntory-gemini-llm-poc"
        self.gcp_location = "us-central1"
        self.embedding_model = "text-embedding-005"
        self._genai_client = genai.Client(
            vertexai=True,
            project=self.gcp_project,
            location=self.gcp_location,
        )

    async def validate_quota(self, user_id: str, new_file_size: int, db: AsyncSession) -> None:
        """Enforce max 5 documents and max 100MB cumulative storage per user."""
        # 1. Check document count
        count = await db.scalar(
            select(func.count(UserDocument.id)).where(UserDocument.user_id == user_id)
        )
        if (count or 0) >= MAX_DOCS_PER_USER:
            raise ValueError(
                f"Document limit reached. You can upload at most {MAX_DOCS_PER_USER} documents."
            )

        # 2. Check cumulative storage
        total_bytes = await db.scalar(
            select(func.coalesce(func.sum(UserDocument.file_size_bytes), 0))
            .where(UserDocument.user_id == user_id)
        )
        total_bytes = total_bytes or 0
        if (total_bytes + new_file_size) > MAX_STORAGE_BYTES:
            remaining_mb = max(0.0, (MAX_STORAGE_BYTES - total_bytes) / (1024 * 1024))
            raise ValueError(
                f"Storage quota exceeded. You have {remaining_mb:.1f} MB remaining out of 100 MB."
            )

    def extract_text(self, filename: str, content: bytes, mime_type: str = "") -> str:
        """Extract clean text from PDF, DOCX, or text/markdown formats."""
        lower_name = filename.lower()
        if lower_name.endswith(".pdf") or "pdf" in mime_type:
            reader = PdfReader(io.BytesIO(content))
            pages_text = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages_text.append(text)
            return "\n\n".join(pages_text)

        if lower_name.endswith(".docx") or "word" in mime_type:
            doc = docx.Document(io.BytesIO(content))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

        # Default: treat as utf-8 plain text / markdown / code
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1", errors="replace")

    def chunk_text(self, text: str, chunk_size: int = 800, overlap: int = 100) -> List[str]:
        """Split raw text into overlapping semantic passages."""
        clean_text = text.strip()
        if not clean_text:
            return []

        chunks = []
        start = 0
        text_len = len(clean_text)

        while start < text_len:
            end = start + chunk_size
            if end >= text_len:
                chunk = clean_text[start:]
                if chunk.strip():
                    chunks.append(chunk.strip())
                break

            # Find convenient split point near whitespace or newline
            split_at = clean_text.rfind("\n\n", start, end)
            if split_at == -1:
                split_at = clean_text.rfind("\n", start, end)
            if split_at == -1:
                split_at = clean_text.rfind(" ", start, end)
            if split_at == -1 or split_at <= start:
                split_at = end

            chunk = clean_text[start:split_at].strip()
            if chunk:
                chunks.append(chunk)
            start = split_at - overlap if split_at - overlap > start else split_at

        return chunks

    async def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate 768-dim embeddings in batches using Vertex AI text-embedding-005."""
        if not texts:
            return []

        embeddings = []
        batch_size = 10
        for i in range(0, len(texts), batch_size):
            batch = texts[i : i + batch_size]
            response = await self._genai_client.aio.models.embed_content(
                model=self.embedding_model,
                contents=batch,
            )
            if hasattr(response, "embeddings"):
                for emb in response.embeddings:
                    embeddings.append(emb.values)
            elif hasattr(response, "embedding"):
                embeddings.append(response.embedding.values)

        return embeddings

    async def ingest_document(
        self,
        user_id: str,
        filename: str,
        content: bytes,
        mime_type: str,
        db: AsyncSession,
    ) -> UserDocument:
        """Parse, validate quota, chunk, embed, and store document in Cloud SQL."""
        file_size = len(content)
        await self.validate_quota(user_id, file_size, db)

        # 1. Create document record with 'indexing' status
        doc = UserDocument(
            user_id=user_id,
            filename=filename,
            file_size_bytes=file_size,
            mime_type=mime_type or "application/octet-stream",
            status="indexing",
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)

        try:
            # 2. Extract and chunk text
            raw_text = self.extract_text(filename, content, mime_type)
            if not raw_text.strip():
                raise ValueError("No extractable text found in uploaded document.")

            chunks = self.chunk_text(raw_text)
            if not chunks:
                raise ValueError("Document yielded 0 text chunks.")

            # 3. Generate embeddings
            vectors = await self.generate_embeddings(chunks)

            # 4. Insert chunks into document_chunks
            for idx, (chunk_text, vec) in enumerate(zip(chunks, vectors)):
                chunk_record = DocumentChunk(
                    document_id=doc.id,
                    chunk_index=idx,
                    chunk_text=chunk_text,
                    token_count=len(chunk_text.split()),
                    embedding=vec,
                )
                db.add(chunk_record)

            doc.status = "ready"
            await db.commit()
            await db.refresh(doc)
            return doc
        except Exception as e:
            logger.error("Failed to ingest document %s: %s", filename, e)
            doc.status = "failed"
            doc.error_message = str(e)
            await db.commit()
            raise

    async def list_documents(self, user_id: str, db: AsyncSession) -> dict[str, Any]:
        """List user documents and compute quota metrics."""
        query = select(UserDocument).order_by(UserDocument.created_at.desc())
        if user_id and user_id != "admin@example.com":
            query = query.where(UserDocument.user_id == user_id)
        result = await db.scalars(query)
        docs = result.all()

        total_bytes = sum(d.file_size_bytes for d in docs)
        return {
            "documents": [
                {
                    "id": str(d.id),
                    "filename": d.filename,
                    "file_size_bytes": d.file_size_bytes,
                    "file_size_mb": round(d.file_size_bytes / (1024 * 1024), 2),
                    "mime_type": d.mime_type,
                    "status": d.status,
                    "error_message": d.error_message,
                    "created_at": d.created_at.isoformat() if d.created_at else None,
                }
                for d in docs
            ],
            "quota": {
                "total_documents": len(docs),
                "max_documents": MAX_DOCS_PER_USER,
                "total_bytes": total_bytes,
                "total_mb": round(total_bytes / (1024 * 1024), 2),
                "max_mb": 100.0,
                "remaining_mb": round(max(0.0, (MAX_STORAGE_BYTES - total_bytes) / (1024 * 1024)), 2),
            },
        }

    async def delete_document(self, doc_id: str, user_id: str, db: AsyncSession) -> bool:
        """Delete a document and cascade its vectorized chunks."""
        doc_uuid = uuid.UUID(doc_id)
        query = select(UserDocument).where(UserDocument.id == doc_uuid)
        if user_id and user_id != "admin@example.com":
            query = query.where(UserDocument.user_id == user_id)
        doc = await db.scalar(query)
        if not doc:
            return False

        await db.execute(delete(DocumentChunk).where(DocumentChunk.document_id == doc_uuid))
        await db.delete(doc)
        await db.commit()
        return True

    async def search_documents(
        self,
        query: str,
        document_ids: List[str],
        top_k: int = 5,
        db: Optional[AsyncSession] = None,
    ) -> List[dict[str, Any]]:
        """Perform cosine similarity search against selected document IDs."""
        if not document_ids or not query.strip() or not db:
            return []

        # 1. Embed query
        query_embeddings = await self.generate_embeddings([query])
        if not query_embeddings:
            return []
        query_vector = query_embeddings[0]

        uuids = [uuid.UUID(d) for d in document_ids]

        # 2. Vector search ordered by cosine distance (<=>)
        stmt = (
            select(
                DocumentChunk.chunk_text,
                DocumentChunk.chunk_index,
                DocumentChunk.document_id,
                UserDocument.filename,
                (1.0 - DocumentChunk.embedding.cosine_distance(query_vector)).label("similarity"),
            )
            .join(UserDocument, DocumentChunk.document_id == UserDocument.id)
            .where(DocumentChunk.document_id.in_(uuids))
            .order_by(DocumentChunk.embedding.cosine_distance(query_vector))
            .limit(top_k)
        )

        result = await db.execute(stmt)
        rows = result.all()

        return [
            {
                "chunk_text": row[0],
                "chunk_index": row[1],
                "document_id": str(row[2]),
                "filename": row[3],
                "similarity": round(float(row[4]), 4),
            }
            for row in rows
        ]


rag_service = RAGService()
